"""阶段七 DOCX 预检与解析契约测试。"""

from hashlib import sha256
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import msoffcrypto
import pytest
from docx import Document
from docx.oxml import OxmlElement

from app.parsing.models import DOCX_MEDIA_TYPE, DocumentParseError, ParseLimits
from app.parsing.normalize import parse_document, stage_upload


def build_docx() -> bytes:
    output = BytesIO()
    document = Document()
    document.add_paragraph("First\u00a0paragraph")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Left cell"
    table.cell(0, 1).text = "Right cell"
    document.add_paragraph("Final paragraph")
    document.save(output)
    return output.getvalue()


def build_blank_docx() -> bytes:
    output = BytesIO()
    Document().save(output)
    return output.getvalue()


def build_encrypted_docx() -> bytes:
    source = BytesIO(build_docx())
    encrypted = BytesIO()
    msoffcrypto.OfficeFile(source).encrypt("test-password", encrypted)
    return encrypted.getvalue()


def build_docx_with_header() -> bytes:
    output = BytesIO()
    document = Document()
    document.add_paragraph("Body text")
    document.sections[0].header.paragraphs[0].text = "Header text"
    document.save(output)
    return output.getvalue()


def build_docx_with_content_control() -> bytes:
    output = BytesIO()
    document = Document()
    controlled_paragraph = document.add_paragraph("Inside content control")
    document.element.body.remove(controlled_paragraph._element)
    content_control = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    content.append(controlled_paragraph._element)
    content_control.append(content)
    document.element.body.insert(0, content_control)
    document.add_paragraph("Outside content control")
    document.save(output)
    return output.getvalue()


def build_forged_docx_package() -> bytes:
    output = BytesIO()
    with ZipFile(output, "w", ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"/>',
        )
        archive.writestr(
            "_rels/.rels",
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>',
        )
        archive.writestr(
            "word/document.xml",
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>',
        )
    return output.getvalue()


def build_docx_with_extra_entry(name: str, content: bytes) -> bytes:
    output = BytesIO()
    with ZipFile(BytesIO(build_docx())) as source, ZipFile(output, "w", ZIP_DEFLATED) as target:
        for entry in source.infolist():
            target.writestr(entry, source.read(entry))
        target.writestr(name, content)
    return output.getvalue()


def test_docx_upload_is_detected_hashed_and_parsed_in_body_order(tmp_path: Path) -> None:
    content = build_docx()

    with stage_upload(
        BytesIO(content),
        original_filename="essay.docx",
        client_media_type=DOCX_MEDIA_TYPE,
        temporary_directory=tmp_path,
        limits=ParseLimits(),
    ) as staged:
        parsed = parse_document(staged, ParseLimits())

        assert staged.media_type == DOCX_MEDIA_TYPE
        assert staged.content_sha256 == sha256(content).digest()
        assert [block.block_id for block in parsed.blocks] == [
            "b000001",
            "b000002",
            "b000003",
            "b000004",
        ]
        assert [block.text for block in parsed.blocks] == [
            "First paragraph",
            "Left cell",
            "Right cell",
            "Final paragraph",
        ]
        assert parsed.blocks[1].locator.model_dump() == {
            "kind": "docx_table_paragraph",
            "table": 1,
            "row": 1,
            "column": 1,
            "paragraph": 1,
        }


def test_encrypted_docx_is_rejected_explicitly(tmp_path: Path) -> None:
    with (
        pytest.raises(DocumentParseError) as error,
        stage_upload(
            BytesIO(build_encrypted_docx()),
            original_filename="encrypted.docx",
            client_media_type=DOCX_MEDIA_TYPE,
            temporary_directory=tmp_path,
            limits=ParseLimits(),
        ),
    ):
        pass

    assert error.value.code == "document_encrypted"


def test_blank_docx_is_rejected_explicitly(tmp_path: Path) -> None:
    with (
        stage_upload(
            BytesIO(build_blank_docx()),
            original_filename="blank.docx",
            client_media_type=DOCX_MEDIA_TYPE,
            temporary_directory=tmp_path,
            limits=ParseLimits(),
        ) as staged,
        pytest.raises(DocumentParseError) as error,
    ):
        parse_document(staged, ParseLimits())

    assert error.value.code == "document_empty"


def test_docx_with_non_body_text_is_rejected_instead_of_silently_omitted(tmp_path: Path) -> None:
    with (
        stage_upload(
            BytesIO(build_docx_with_header()),
            original_filename="header.docx",
            client_media_type=DOCX_MEDIA_TYPE,
            temporary_directory=tmp_path,
            limits=ParseLimits(),
        ) as staged,
        pytest.raises(DocumentParseError) as error,
    ):
        parse_document(staged, ParseLimits())

    assert error.value.code == "docx_content_unsupported"


def test_docx_content_controls_are_rejected_instead_of_silently_omitted(tmp_path: Path) -> None:
    with (
        stage_upload(
            BytesIO(build_docx_with_content_control()),
            original_filename="content-control.docx",
            client_media_type=DOCX_MEDIA_TYPE,
            temporary_directory=tmp_path,
            limits=ParseLimits(),
        ) as staged,
        pytest.raises(DocumentParseError) as error,
    ):
        parse_document(staged, ParseLimits())

    assert error.value.code == "docx_content_unsupported"


def test_zip_with_docx_filenames_but_no_ooxml_contract_is_rejected(tmp_path: Path) -> None:
    with (
        stage_upload(
            BytesIO(build_forged_docx_package()),
            original_filename="forged.docx",
            client_media_type=DOCX_MEDIA_TYPE,
            temporary_directory=tmp_path,
            limits=ParseLimits(),
        ) as staged,
        pytest.raises(DocumentParseError) as error,
    ):
        parse_document(staged, ParseLimits())

    assert error.value.code == "docx_archive_invalid"


def test_docx_uncompressed_size_limit_is_enforced_before_parsing(tmp_path: Path) -> None:
    with (
        stage_upload(
            BytesIO(build_docx()),
            original_filename="large-expanded.docx",
            client_media_type=DOCX_MEDIA_TYPE,
            temporary_directory=tmp_path,
            limits=ParseLimits(),
        ) as staged,
        pytest.raises(DocumentParseError) as error,
    ):
        parse_document(staged, ParseLimits(max_uncompressed_bytes=1_000))

    assert error.value.code == "docx_archive_too_large"


@pytest.mark.parametrize(
    ("filename", "content", "expected_code"),
    [
        ("../escape.xml", b"<root/>", "docx_archive_invalid"),
        ("word/vbaProject.bin", b"macro", "docx_macro_unsupported"),
        ("word/unsafe.xml", b"<!DOCTYPE root><root/>", "docx_xml_unsafe"),
        ("word/document.xml", b"<duplicate/>", "docx_archive_invalid"),
    ],
)
@pytest.mark.filterwarnings("ignore:Duplicate name")
def test_docx_archive_rejects_unsafe_entries(
    tmp_path: Path,
    filename: str,
    content: bytes,
    expected_code: str,
) -> None:
    with (
        stage_upload(
            BytesIO(build_docx_with_extra_entry(filename, content)),
            original_filename="unsafe.docx",
            client_media_type=DOCX_MEDIA_TYPE,
            temporary_directory=tmp_path,
            limits=ParseLimits(),
        ) as staged,
        pytest.raises(DocumentParseError) as error,
    ):
        parse_document(staged, ParseLimits())

    assert error.value.code == expected_code


def test_docx_compression_ratio_limit_is_enforced_before_parsing(tmp_path: Path) -> None:
    with (
        stage_upload(
            BytesIO(build_docx()),
            original_filename="compressed.docx",
            client_media_type=DOCX_MEDIA_TYPE,
            temporary_directory=tmp_path,
            limits=ParseLimits(),
        ) as staged,
        pytest.raises(DocumentParseError) as error,
    ):
        parse_document(staged, ParseLimits(max_zip_compression_ratio=1))

    assert error.value.code == "docx_archive_too_large"
