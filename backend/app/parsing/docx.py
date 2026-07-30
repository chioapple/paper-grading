"""DOCX 结构验证和顺序化文本块解析。"""

import zipfile
from pathlib import PurePosixPath
from xml.etree import ElementTree

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.parsing.models import (
    DOCX_MEDIA_TYPE,
    DocumentBlock,
    DocumentParseError,
    DocxParagraphLocator,
    DocxTableParagraphLocator,
    ParsedDocument,
    ParseLimits,
    StagedDocument,
)
from app.parsing.normalize import DOCX_REQUIRED_ENTRIES, normalize_text

CONTENT_TYPES_NAMESPACE = "http://schemas.openxmlformats.org/package/2006/content-types"
RELATIONSHIPS_NAMESPACE = "http://schemas.openxmlformats.org/package/2006/relationships"
DOCX_MAIN_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
)
OFFICE_DOCUMENT_RELATIONSHIP = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument"
)
UNSAFE_ENTRY_PREFIXES = ("word/embeddings/", "word/activex/")
UNSAFE_RELATIONSHIP_TYPES = {
    "http://schemas.microsoft.com/office/2006/relationships/activexcontrol",
    "http://schemas.microsoft.com/office/2006/relationships/activexcontrolbinary",
    "http://schemas.openxmlformats.org/officedocument/2006/relationships/control",
    "http://schemas.openxmlformats.org/officedocument/2006/relationships/oleobject",
    "http://schemas.openxmlformats.org/officedocument/2006/relationships/package",
}
MACRO_RELATIONSHIP_TYPES = {
    "http://schemas.microsoft.com/office/2006/relationships/vbaproject",
    "http://schemas.openxmlformats.org/officedocument/2006/relationships/vbaproject",
}


def _validate_archive(document: StagedDocument, limits: ParseLimits) -> None:
    try:
        with zipfile.ZipFile(document.path) as archive:
            entries = archive.infolist()
            names = [entry.filename for entry in entries]
            if len(entries) > limits.max_zip_entries or len(names) != len(set(names)):
                raise DocumentParseError("docx_archive_invalid", "DOCX 压缩结构无效")
            if not set(names) >= DOCX_REQUIRED_ENTRIES:
                raise DocumentParseError("docx_archive_invalid", "DOCX 缺少必要结构")
            total_uncompressed = 0
            for entry in entries:
                path = PurePosixPath(entry.filename)
                normalized_name = entry.filename.casefold()
                if path.is_absolute() or ".." in path.parts or entry.flag_bits & 0x1:
                    raise DocumentParseError("docx_archive_invalid", "DOCX 压缩结构无效")
                if normalized_name == "word/vbaproject.bin":
                    raise DocumentParseError(
                        "docx_macro_unsupported",
                        "不支持包含宏的 Word 文件",
                    )
                if normalized_name.startswith(UNSAFE_ENTRY_PREFIXES):
                    raise DocumentParseError(
                        "docx_active_content_unsupported",
                        "不支持包含内嵌对象或 ActiveX 的 Word 文件",
                    )
                total_uncompressed += entry.file_size
                if total_uncompressed > limits.max_uncompressed_bytes:
                    raise DocumentParseError("docx_archive_too_large", "DOCX 解压后内容过大")
                ratio: float
                if entry.compress_size == 0:
                    ratio = float(entry.file_size)
                else:
                    ratio = entry.file_size / entry.compress_size
                if ratio > limits.max_zip_compression_ratio:
                    raise DocumentParseError("docx_archive_too_large", "DOCX 压缩比例异常")
                if entry.filename.endswith((".xml", ".rels")):
                    xml = archive.read(entry)
                    if b"<!DOCTYPE" in xml.upper() or b"<!ENTITY" in xml.upper():
                        raise DocumentParseError("docx_xml_unsafe", "DOCX 包含不安全 XML")
                    if entry.filename.endswith(".rels"):
                        relationship_tree = ElementTree.fromstring(xml)
                        for item in relationship_tree.findall(
                            f"{{{RELATIONSHIPS_NAMESPACE}}}Relationship"
                        ):
                            if item.get("TargetMode") == "External":
                                raise DocumentParseError(
                                    "docx_external_relationship",
                                    "DOCX 包含外部关系",
                                )
                            relationship_type = (item.get("Type") or "").casefold()
                            if relationship_type in MACRO_RELATIONSHIP_TYPES:
                                raise DocumentParseError(
                                    "docx_macro_unsupported",
                                    "不支持包含宏的 Word 文件",
                                )
                            if relationship_type in UNSAFE_RELATIONSHIP_TYPES:
                                raise DocumentParseError(
                                    "docx_active_content_unsupported",
                                    "不支持包含内嵌对象或 ActiveX 的 Word 文件",
                                )
            content_types = ElementTree.fromstring(archive.read("[Content_Types].xml"))
            has_document_content_type = any(
                item.get("PartName") == "/word/document.xml"
                and item.get("ContentType") == DOCX_MAIN_CONTENT_TYPE
                for item in content_types.findall(f"{{{CONTENT_TYPES_NAMESPACE}}}Override")
            )
            relationships = ElementTree.fromstring(archive.read("_rels/.rels"))
            has_document_relationship = any(
                item.get("Type") == OFFICE_DOCUMENT_RELATIONSHIP
                and (item.get("Target") or "").lstrip("/") == "word/document.xml"
                and item.get("TargetMode") != "External"
                for item in relationships.findall(f"{{{RELATIONSHIPS_NAMESPACE}}}Relationship")
            )
            if not has_document_content_type or not has_document_relationship:
                raise DocumentParseError("docx_archive_invalid", "DOCX OOXML 主文档契约无效")
            damaged = archive.testzip()
            if damaged is not None:
                raise DocumentParseError("docx_archive_invalid", "DOCX 压缩内容损坏")
    except DocumentParseError:
        raise
    except (
        ElementTree.ParseError,
        OSError,
        RuntimeError,
        zipfile.BadZipFile,
        zipfile.LargeZipFile,
    ) as error:
        raise DocumentParseError("docx_archive_invalid", "DOCX 文件损坏") from error


def parse_docx(document: StagedDocument, limits: ParseLimits) -> ParsedDocument:
    """按正文顺序提取段落和表格单元格段落。"""

    _validate_archive(document, limits)
    try:
        source = Document(str(document.path))
    except (KeyError, OSError, ValueError, zipfile.BadZipFile) as error:
        raise DocumentParseError("docx_parse_failed", "DOCX 内容无法解析") from error

    if source.element.body.xpath(".//w:txbxContent | .//w:ins | .//w:del | .//w:sdt"):
        raise DocumentParseError(
            "docx_content_unsupported",
            "DOCX 包含文本框、内容控件或未接受的修订内容",
        )
    for section in source.sections:
        stories = (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        )
        for story in stories:
            if any(normalize_text(paragraph.text) for paragraph in story.paragraphs):
                raise DocumentParseError(
                    "docx_content_unsupported",
                    "DOCX 页眉或页脚包含首版不支持的文字",
                )
            if any(
                normalize_text(cell.text)
                for table in story.tables
                for row in table.rows
                for cell in row.cells
            ):
                raise DocumentParseError(
                    "docx_content_unsupported",
                    "DOCX 页眉或页脚包含首版不支持的表格文字",
                )
            if story._element.xpath(".//w:txbxContent//w:t"):
                raise DocumentParseError(
                    "docx_content_unsupported",
                    "DOCX 页眉或页脚包含首版不支持的文本框",
                )

    blocks: list[DocumentBlock] = []
    character_count = 0

    def append_block(text: str, locator: DocxParagraphLocator | DocxTableParagraphLocator) -> None:
        nonlocal character_count
        normalized = normalize_text(text)
        if not normalized:
            return
        character_count += len(normalized)
        if character_count > limits.max_characters:
            raise DocumentParseError("document_text_too_large", "可提取文本超过限制")
        if len(blocks) >= limits.max_blocks:
            raise DocumentParseError("document_blocks_too_many", "可定位文本块超过限制")
        blocks.append(
            DocumentBlock(
                block_id=f"b{len(blocks) + 1:06d}",
                text=normalized,
                locator=locator,
            )
        )

    paragraph_number = 0
    table_number = 0
    for item in source.iter_inner_content():
        if isinstance(item, Paragraph):
            paragraph_number += 1
            append_block(item.text, DocxParagraphLocator(paragraph=paragraph_number))
            continue
        if not isinstance(item, Table):
            raise DocumentParseError("docx_content_unsupported", "DOCX 包含不支持的正文结构")
        table_number += 1
        seen_cells: set[int] = set()
        for row_number, row in enumerate(item.rows, start=1):
            for column_number, cell in enumerate(row.cells, start=1):
                cell_identity = id(cell._tc)
                if cell_identity in seen_cells:
                    continue
                seen_cells.add(cell_identity)
                if cell.tables:
                    raise DocumentParseError("docx_content_unsupported", "不支持嵌套表格")
                for cell_paragraph_number, paragraph in enumerate(cell.paragraphs, start=1):
                    append_block(
                        paragraph.text,
                        DocxTableParagraphLocator(
                            table=table_number,
                            row=row_number,
                            column=column_number,
                            paragraph=cell_paragraph_number,
                        ),
                    )

    if not blocks:
        raise DocumentParseError("document_empty", "文档没有可提取文字")
    return ParsedDocument(
        media_type=DOCX_MEDIA_TYPE,
        page_count=None,
        character_count=character_count,
        blocks=blocks,
    )
